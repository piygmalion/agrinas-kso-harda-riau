# -*- coding: utf-8 -*-
"""Generate laporan analisis konflik agraria — Unit II Harda Ditreskrimum Polda Riau."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from copy import deepcopy


OUT = r"C:\Users\Patron\Downloads\sawit lagi\Laporan_Analisis_Konflik_Agraria_Agrinas_KSO_Polda_Riau.docx"

NAVY = RGBColor(0x1A, 0x3A, 0x5C)
DARK = RGBColor(0x22, 0x22, 0x22)
GRAY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "1A3A5C"
ALT_ROW = "F0F4F8"
RED_BG = "C0392B"
YELLOW_BG = "F39C12"
GREEN_BG = "27AE60"


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=DARK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell, color="666666", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_text(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, color=DARK, font="Times New Roman"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(str(text))
    set_run_font(run, name=font, size=size, bold=bold, color=color)


def add_para(doc, text, size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_after=6, space_before=0, first_indent=None, color=DARK):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.15
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_rich_para(doc, parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6, first_indent=None):
    """parts: list of (text, bold, italic, size, color)"""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if first_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    for part in parts:
        text = part[0]
        bold = part[1] if len(part) > 1 else False
        italic = part[2] if len(part) > 2 else False
        size = part[3] if len(part) > 3 else 12
        color = part[4] if len(part) > 4 else DARK
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading_custom(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    if level == 1:
        pf.space_before = Pt(18)
        pf.space_after = Pt(8)
        size, bold = 14, True
        # bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "1A3A5C")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 2:
        pf.space_before = Pt(14)
        pf.space_after = Pt(6)
        size, bold = 12, True
    else:
        pf.space_before = Pt(10)
        pf.space_after = Pt(4)
        size, bold = 12, True
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=NAVY)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    if level:
        p.paragraph_format.left_indent = Cm(1.25 + level * 0.5)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True)
        r2 = p.add_run(text)
        set_run_font(r2, size=11)
    else:
        # clear default empty run if any
        if p.runs:
            p.runs[0].text = ""
        run = p.add_run(text)
        set_run_font(run, size=11)
    return p


def add_numbered(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=11, bold=True)
        r2 = p.add_run(text)
        set_run_font(r2, size=11)
    else:
        if p.runs:
            p.runs[0].text = ""
        run = p.add_run(text)
        set_run_font(run, size=11)
    return p


def make_table(doc, headers, rows, col_widths=None, header_bg=HEADER_BG):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=WHITE)
        shade_cell(cell, header_bg)
        set_cell_borders(cell, color="1A3A5C", sz="4")

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 and len(headers) > 3 else WD_ALIGN_PARAGRAPH.LEFT
            # numeric-looking columns center
            if isinstance(val, (int, float)) or (isinstance(val, str) and val.replace(".", "").replace(",", "").replace("*", "").isdigit()):
                align = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cell, val, bold=False, size=9, align=align)
            set_cell_borders(cell, color="888888", sz="4")
            if r_idx % 2 == 1:
                shade_cell(cell, ALT_ROW)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r1 = p.add_run(label)
    set_run_font(r1, size=11, bold=True, color=NAVY)
    r2 = p.add_run(value)
    set_run_font(r2, size=11)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)
    # shade via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F5F5")
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)
    run = p.add_run(text)
    set_run_font(run, name="Consolas", size=9, color=DARK)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hp.add_run("UNIT II HARDA — DITRESKRIMUM POLDA RIAU")
    set_run_font(r, size=9, bold=True, color=NAVY)
    hp2 = header.add_paragraph()
    hp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = hp2.add_run("Laporan Analisis Konflik Agraria Agrinas–KSO | Rahasia Internal — Analisis")
    set_run_font(r2, size=8, italic=True, color=GRAY)
    # header border
    pPr = hp2._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1A3A5C")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Unit II Harda Ditreskrimum Polda Riau  |  Halaman ")
    set_run_font(r, size=8, color=GRAY)
    add_page_number(fp)
    r3 = fp.add_run("  |  3 Agustus 2026")
    set_run_font(r3, size=8, color=GRAY)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    return doc


def build():
    doc = setup_document()

    # ═══ COVER ═══
    for _ in range(2):
        doc.add_paragraph()

    p = add_para(doc, "KEPOLISIAN NEGARA REPUBLIK INDONESIA", size=12, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "DAERAH RIAU", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "DIREKTORAT RESERSE KRIMINAL UMUM", size=11, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "UNIT II HARDA", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # line
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = line._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A3A5C")
    pBdr.append(bottom)
    pPr.append(pBdr)

    add_para(doc, "LAPORAN ANALISIS", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=16, space_after=6, color=NAVY)
    add_para(doc, "KONFLIK AGRARIA AGRINAS–KSO\nDI 12 POLRES WILAYAH HUKUM POLDA RIAU",
             size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16, color=NAVY)

    add_para(doc, "Analisis Berbasis Dokumen Polri / Intelkam / Reskrim",
             size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20, color=GRAY)

    # meta box as table
    meta = [
        ("Unit Penyusun", "Unit II Harda — Ditreskrimum Polda Riau"),
        ("Periode Sumber", "Terutama 2024 – Juli 2026 (log Kuansing hingga 2020)"),
        ("Tanggal Penyusunan", "3 Agustus 2026"),
        ("Sifat Dokumen", "Deskripsi–sintesis dokumen; bukan verifikasi lapangan"),
        ("Produk Pendukung", "matriks_agrinas_kso_12_polres.xlsx"),
        ("Klasifikasi", "Internal — untuk kepentingan analisis Unit II Harda"),
    ]
    t = doc.add_table(rows=len(meta), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (lab, val) in enumerate(meta):
        set_cell_text(t.rows[i].cells[0], lab, bold=True, size=10, color=WHITE)
        shade_cell(t.rows[i].cells[0], HEADER_BG)
        set_cell_borders(t.rows[i].cells[0])
        set_cell_text(t.rows[i].cells[1], val, size=10)
        set_cell_borders(t.rows[i].cells[1])
        t.rows[i].cells[0].width = Cm(5)
        t.rows[i].cells[1].width = Cm(11)

    doc.add_paragraph()
    add_para(doc, "Pekanbaru, Agustus 2026", size=11, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=30, color=GRAY)

    doc.add_page_break()

    # ═══ DAFTAR ISI (manual singkat) ═══
    add_heading_custom(doc, "DAFTAR ISI", level=1)
    toc_items = [
        "1. Ringkasan Eksekutif",
        "2. Metode dan Batasan",
        "3. Profil Tingkat Polda (Baseline)",
        "4. Analisis Bertahap per Polres",
        "5. Sintesis Lintas 12 Polres",
        "6. Rekomendasi Analisis Lanjutan",
        "7. Lampiran Sumber",
    ]
    for item in toc_items:
        add_para(doc, item, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

    doc.add_page_break()

    # ═══ 1. RINGKASAN EKSEKUTIF ═══
    add_heading_custom(doc, "1. RINGKASAN EKSEKUTIF", level=1)

    add_para(doc,
        "Ekosistem pengelolaan lahan sawit sitaan Satgas PKH di Riau melibatkan 130 perusahaan "
        "dalam skema PT Agrinas Palma Nusantara dan mitra KSO, dengan klaster Intelkam "
        "3 merah · 24 kuning · 103 hijau. Konflik berdarah dan kerusakan material terkonsentrasi "
        "di koridor utara (Rohul–Rohil–Bengkalis), sementara volume kebun terbesar ada di Inhu. "
        "Banyak kebun berlabel hijau tetap menunjukkan gesekan di sumber satker (Inhu, Kuansing, Kampar)— "
        "menandakan klaster Intelkam perlu divalidasi ulang terhadap kejadian aktual 2025–2026.",
        first_indent=1.0)

    add_para(doc, "Temuan inti untuk Unit II Harda:", bold=True, space_before=8, space_after=4,
             align=WD_ALIGN_PARAGRAPH.LEFT)

    findings = [
        "Konflik struktural tata kelola (penunjukan KSO, PAM non-BUJP, batas pasca-PKH, klaim plasma/MHA) "
        "lebih determinan terhadap eskalasi daripada pencurian TBS oportunistik semata.",
        "Tiga titik merah resmi: Gunung Mas Raya–UTS (Rohil), Berkat Satu–Majuma (Rohul, 1 MD), "
        "SIS–PAB (Bengkalis).",
        "Pelalawan khusus: daftar KSO Intelkam seluruhnya hijau, tetapi dimensi TNTN menghasilkan "
        "aksi kolektif tinggi—harus dipisah dari analisis KSO murni.",
        "Pekanbaru dan Kep. Meranti hampir NIHIL untuk Agrinas–KSO; jangan digabung dengan pidana "
        "pencurian masyarakat.",
        "Kualitas data tidak merata; angka luas Pelalawan/Siak/sebagian Kuansing tidak layak "
        "diagregasi tanpa klarifikasi.",
    ]
    for i, f in enumerate(findings, 1):
        add_para(doc, f"{i}. {f}", size=11, space_after=4, first_indent=0)

    # ═══ 2. METODE ═══
    add_heading_custom(doc, "2. METODE DAN BATASAN", level=1)

    add_heading_custom(doc, "2.1 Kerangka Analitik (Seragam Tiap Polres)", level=2)
    kerangka = [
        "Inventaris lahan (eks-perusahaan, luas sita/PKH, status KSO, peran Agrinas)",
        "Tipologi konflik (aktual vs potensi)",
        "Peta aktor",
        "Intensitas & dampak (LP, korban, klaster)",
        "Respons tercatat",
        "Kualitas bukti (kuat / sedang / tipis / NIHIL)",
    ]
    for i, k in enumerate(kerangka, 1):
        add_para(doc, f"{i}. {k}", size=11, space_after=2, align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading_custom(doc, "2.2 Sumber Utama", level=2)
    make_table(doc,
        ["Sumber", "Peran"],
        [
            ["Intelkam 12.02.2026 — Permasalahan KSO Agrinas", "Baseline 130 perusahaan + klaster + bentrok"],
            ["Intelkam 23.02.2026 — TNTN & KSO", "Update TNTN Pelalawan + kasus Majuma"],
            ["Harda 2026 — Permasalahan Tanah", "Framing tipologi yuridis/sosial/administratif/keamanan"],
            ["Laporan satker (Kampar, Inhu, Rohul, Bengkalis, Kuansing, Inhil, Siak, Pelalawan, dll.)",
             "Detail lokal LP/kebun"],
        ],
        col_widths=[8.5, 8])

    add_heading_custom(doc, "2.3 Batasan", level=2)
    batasans = [
        "Tidak ada dataset CSV master; ekstraksi dari PDF/DOCX dengan noise OCR.",
        "Filename menyesatkan (mis. “POLDA RIAU” berisi hanya Kampar).",
        "Beberapa satker hanya agregat atau NIHIL.",
        "Analisis tidak memberikan rekomendasi taktis operasi; fokus rekomendasi analisis lanjutan.",
    ]
    for b in batasans:
        add_bullet(doc, b)

    # ═══ 3. PROFIL POLDA ═══
    add_heading_custom(doc, "3. PROFIL TINGKAT POLDA (BASELINE)", level=1)

    add_heading_custom(doc, "3.1 Skala Ekosistem Agrinas–KSO", level=2)
    make_table(doc,
        ["Satker", "Jml Kebun", "Merah", "Kuning", "Hijau", "Konflik (resume)"],
        [
            ["Inhu", "28", "0", "2", "26", "3"],
            ["Rohil", "24", "1", "6", "17", "7"],
            ["Rohul", "21", "1", "5", "15", "6"],
            ["Inhil", "15", "0", "4", "11", "4"],
            ["Kuansing", "11", "0", "1", "10", "1"],
            ["Pelalawan", "9", "0", "0", "9", "0"],
            ["Bengkalis", "8", "1", "0", "7", "3"],
            ["Siak", "6", "0", "0", "6", "1"],
            ["Kampar", "4", "0", "3", "1", "3"],
            ["Dumai", "3", "0", "3", "0", "3"],
            ["Kep. Meranti", "1", "0", "0", "1", "0"],
            ["Pekanbaru", "0*", "—", "—", "—", "—"],
            ["TOTAL", "130", "3", "24", "103", ""],
        ],
        col_widths=[3.2, 2.2, 1.8, 1.8, 1.8, 3.2])

    add_para(doc,
        "* Pekanbaru tidak masuk breakdown 130; laporan lokal menyatakan Agrinas/KSO NIHIL.",
        size=9, italic=True, color=GRAY, space_after=8)

    add_heading_custom(doc, "3.2 Definisi Klaster Intelkam", level=2)
    add_bullet(doc, "konflik + kerusuhan + korban", bold_prefix="Merah: ")
    add_bullet(doc, "konflik tanpa kerusuhan", bold_prefix="Kuning: ")
    add_bullet(doc, "belum terjadi konflik", bold_prefix="Hijau: ")

    add_heading_custom(doc, "3.3 Kejadian Menonjol Tingkat Provinsi", level=2)
    kejadian = [
        "6 bentrok di Rohil, Rohul, Bengkalis.",
        "Korban: 1 meninggal dunia + 23 luka.",
        "Kerugian material (resume): 13 mobil, 3 motor, 4 rumah, 14 HP.",
        "6 perusahaan disebut memakai PAM swakarsa non-BUJP (antara lain Majuma/Rohul, Torus/Rohul, "
        "PAB/Bengkalis, UTS/Rohil, CV Makmur Jaya Sentosa/Kampar).",
    ]
    for k in kejadian:
        add_bullet(doc, k)

    add_heading_custom(doc, "3.4 Pola Pemicu Generik (Intelkam + Harda)", level=2)
    pemicu = [
        "Tapal batas belum jelas pasca sitaan Satgas PKH.",
        "Penerima KSO bukan penguasa lahan / bukan warga tempatan.",
        "PAM swakarsa non-BUJP, isu SARA, premanisme.",
        "Minim koordinasi Pemda / APH.",
        "Sejarah konflik lama diangkat kembali.",
        "Klaim masyarakat / plasma terputus akibat penertiban.",
        "Tidak ada standar seragam Agrinas untuk penunjukan KSO.",
        "Tipologi Harda: yuridis (HGU/ulayat/hutan), sosial (adat, kesenjangan), "
        "administratif (peta sektoral), keamanan (okupasi massal, ormas).",
    ]
    for i, p in enumerate(pemicu, 1):
        add_para(doc, f"{i}. {p}", size=11, space_after=2, align=WD_ALIGN_PARAGRAPH.LEFT)

    add_heading_custom(doc, "3.5 Ranking Prioritas Analitik Harda", level=2)
    make_table(doc,
        ["Prioritas", "Satker", "Alasan Singkat"],
        [
            ["Sangat tinggi", "Rohul, Rohil, Bengkalis", "Merah + bentrok + korban"],
            ["Tinggi", "Inhu, Kuansing, Kampar, Dumai, Pelalawan (TNTN)", "Volume / LP / dimensi kawasan"],
            ["Sedang", "Inhil, Siak", "Konflik ada; data parsial"],
            ["Rendah (gap-fill)", "Kep. Meranti, Pekanbaru", "NIHIL Agrinas–KSO"],
        ],
        col_widths=[3.5, 7, 6])

    doc.add_page_break()

    # ═══ 4. ANALISIS PER POLRES ═══
    add_heading_custom(doc, "4. ANALISIS BERTAHAP PER POLRES", level=1)
    add_para(doc,
        "Urutan mengikuti beban kasus/kebun dan kekayaan sumber (bukan alfabet).",
        italic=True, size=11, color=GRAY)

    # --- 4.1 Inhu ---
    add_heading_custom(doc, "4.1 Polres Indragiri Hulu (Inhu)", level=2)
    add_rich_para(doc, [
        ("Ringkasan satker. ", True),
        ("Volume kebun Agrinas/KSO tertinggi (", False),
        ("28", True),
        ("). Klaster resmi 0 merah / 2 kuning / 26 hijau, tetapi sumber lokal memuat banyak "
         "gesekan plasma, MHA, dan panen paksa yang belum sepenuhnya tercermin sebagai “merah”.", False),
    ])
    add_rich_para(doc, [
        ("Inventaris (cuplikan). ", True),
        ("Bernas Mulya Mandiri mengelola beberapa unit besar BBU/KAT (ribuan Ha); "
         "PT Indrawan Perkasa–Tiga Raja Mas (kuning, bentrok pemanen); "
         "PT Selantai Argo Lestari (kuning, MHA Talang Mamak); "
         "EX Palm Lestari–Koperasi TKBM (sengketa 560 Ha, panen paksa Apr 2026); "
         "TPP–Poktan JD Karya Mandiri (aksi Jun 2026); sejumlah unit dikelola langsung Agrinas "
         "(Seberida Subur, Duta Palma).", False),
    ])
    add_rich_para(doc, [
        ("Tipologi. ", True),
        ("Dominan: tuntutan plasma/kelola bersama, konflik adat, gesekan antar pemanen, "
         "pencurian TBS sporadis, penolakan sita plasma (TDE/SML).", False),
    ])
    add_rich_para(doc, [
        ("Aktor. ", True),
        ("Agrinas, mitra KSO personal/koperasi/CV, eks-perusahaan, poktan, MHA Talang Mamak, "
         "advokat/kelompok penuntut lahan.", False),
    ])
    add_rich_para(doc, [
        ("Intensitas. ", True),
        ("Resume Intelkam: 3 konflik. LP formal jarang di tabel lokal; narasi konflik relatif kaya.", False),
    ])
    add_rich_para(doc, [
        ("Respons. ", True),
        ("Mediasi (Redang Seko), monitoring Satgas/Agrinas; hambatan: mapping penerima KSO vs warga tempatan.", False),
    ])
    add_rich_para(doc, [
        ("Kualitas bukti. ", True),
        ("Kuat", True),
        (" untuk inventaris naratif; ", False),
        ("sedang", True),
        (" untuk LP bernomor.", False),
    ])
    add_rich_para(doc, [
        ("Flag lanjutan: ", True, True),
        ("validasi klaster hijau vs kejadian lokal; jaringan penerima multi-lokasi "
         "(Bernas Mulya Mandiri, PT Runggu).", False, True),
    ], space_after=8)

    # --- 4.2 Rohil ---
    add_heading_custom(doc, "4.2 Polres Rokan Hilir (Rohil)", level=2)
    add_rich_para(doc, [
        ("Ringkasan. ", True),
        ("24 kebun (1 merah / 6 kuning / 17 hijau); konflik resume ", False),
        ("7", True),
        ("—tertinggi numerik bersama koridor utara.", False),
    ])
    add_para(doc, "Inventaris kunci:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_bullet(doc, "PT Gunung Mas Raya (Rumbia I) – KSO PT Ujung Tanjung Sejahtera (PAM Flores non-BUJP).",
               bold_prefix="MERAH: ")
    add_bullet(doc, "Salim Ivomas–Digjaya Nata Persada (tolak serah KSO); Rama Salomo (2 unit); "
               "APSL–Satahi/Poktan Melayu Terpadu; Bumi Riau Bina Makmur–K21.",
               bold_prefix="KUNING: ")
    add_bullet(doc, "sitaan PKH 27.655,51 Ha = total Agrinas/KSO; 8 KSO (nama tidak lengkap di Bismillah.docx).",
               bold_prefix="Agregat lokal: ")
    add_rich_para(doc, [
        ("Tipologi. ", True),
        ("Penolakan penyerahan/plang KSO; bentrok panen sepihak; PAM non-BUJP.", False),
    ])
    add_rich_para(doc, [
        ("Kasus unggulan. ", True),
        ("20 Okt 2025 — kelompok W. Siringo-Ringo vs security UTS; 7 luka; RJ.", False),
    ])
    add_rich_para(doc, [
        ("Kualitas bukti. ", True),
        ("Intelkam kuat; detail LP Agrinas (agregat=3) tipis. ", False),
        ("Flag: ", True, True),
        ("itemisasi 3 LP Agrinas + daftar 8 KSO lengkap.", False, True),
    ], space_after=8)

    # --- 4.3 Rohul ---
    add_heading_custom(doc, "4.3 Polres Rokan Hulu (Rohul)", level=2)
    add_rich_para(doc, [
        ("Ringkasan. ", True),
        ("21 kebun (1 merah / 5 kuning / 15 hijau). Titik paling kritis secara dampak: ", False),
        ("1 MD", True),
        (".", False),
    ])
    add_para(doc, "Inventaris kunci:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_bullet(doc, "PT BK1/Berkat Satu – KSO Nusantara Sawit Majuma (Bonai Darussalam).",
               bold_prefix="MERAH: ")
    add_bullet(doc, "Torganda Tambusai Timur; Togos Gopas vs Torus Ganda; Ekaudra–CV Ginting (tolak 5 desa); "
               "Torganda Batang Kumu 2; Torganda Rantau Kasai (~3.000 Ha panen sepihak).",
               bold_prefix="KUNING: ")
    add_bullet(doc, "12 KSO bernama; sitaan PKH 65.947,06 Ha.", bold_prefix="Lokal: ")

    add_para(doc, "Kasus LP (semua P21 di sumber Rohul):", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)
    make_table(doc,
        ["LP", "Pokok", "Dampak"],
        [
            ["LP/B/07/II/2026 Bonai", "Serangan PAM Majuma vs PAM Telago Biru/Berkat Satu", "1 MD + luka"],
            ["LP/B/107/III/2026", "Pencurian TBS Agrinas eks Torganda", "~2 ton"],
            ["LP/B/32/I/2026", "Penganiayaan patroli Agrinas+Satgas", "Luka"],
            ["LP/B/21/II/2026 Tambusai", "Massa ~400 paksa masuk mess Afd VIII", "Kerusakan"],
        ],
        col_widths=[4.5, 8, 4])

    add_rich_para(doc, [
        ("Tipologi. ", True),
        ("Bentrok antar PAM swakarsa; penolakan KSO desa; panen sepihak eks-mitra/eks-karyawan.", False),
    ])
    add_rich_para(doc, [
        ("Kualitas bukti. ", True),
        ("Kuat. ", True),
        ("Catatan: jumlah tersangka Majuma berbeda antara deck 12.02 dan 23.02. ", False),
        ("Flag: ", True, True),
        ("network PAM/KSO; timeline eskalasi Jan–Feb 2026.", False, True),
    ], space_after=8)

    # --- 4.4 Inhil ---
    add_heading_custom(doc, "4.4 Polres Indragiri Hilir (Inhil)", level=2)
    add_rich_para(doc, [
        ("Ringkasan. ", True),
        ("15 kebun (0 merah / 4 kuning / 11 hijau); konflik resume 4.", False),
    ])
    add_rich_para(doc, [
        ("Inventaris/konflik. ", True),
        ("PT Agro Sarimas–Citra Mutiara Bumi Riau (kuning, tuntutan kembalikan ke KPCH); "
         "PT RSA–CV Cipta Nugraha (Mandah, 750,21 Ha) — pemanenan TBS Mei 2026 tahap penyelidikan; "
         "isu penghadangan vs APN/Kemuning; KT Naibaho (pencurian, bukan konflik KSO aktif).", False),
    ])
    add_rich_para(doc, [
        ("Tipologi. ", True),
        ("Perebutan legitimasi kelola pasca-PKH; dugaan panen oleh KSO vs eks-operator.", False),
    ])
    add_rich_para(doc, [
        ("Kualitas bukti. ", True),
        ("Sedang", True),
        (" — Intelkam lebih kaya dari PDF konflik lokal (hanya 2 baris estate). ", False),
        ("Flag: ", True, True),
        ("lengkapi 15 kartu Intelkam ke matriks satker; tarik nomor LP RSA.", False, True),
    ], space_after=8)

    # --- 4.5 Kuansing ---
    add_heading_custom(doc, "4.5 Polres Kuantan Singingi (Kuansing)", level=2)
    add_rich_para(doc, [
        ("Ringkasan. ", True),
        ("11 kebun (0M / 1K / 10H). Pola khas: ", False),
        ("pidana pencurian TBS berulang", True),
        (" di areal yang kemudian/kini dikelola Agrinas, plus potensi struktural di Pesikaian/WJT.", False),
    ])
    add_rich_para(doc, [
        ("Inventaris. ", True),
        ("Cerenti Subur (8.929 Ha, belum KSO, kelola Agrinas); WJT (4.196 Ha); "
         "PTPN IV Afd 7–9 Pesikaian (kuning, relokasi TNTN); "
         "Merauke Tetap Jaya–Wana Agri Santosa (kuning); beberapa unit sudah KSO "
         "(Garuda, Cahaya Panam, Sailan, Gatipura).", False),
    ])
    add_rich_para(doc, [
        ("Kasus. ", True),
        ("Log 2020–2026: ~52 LP di areal kebun; ~35–37 relevan Agrinas-estate "
         "(Cerenti, WJT, PTPN, Agrinas). Potensi 2026: relokasi Pesikaian; tuntutan 20% WJT; "
         "isu koperasi Siampo Pelangi.", False),
    ])
    add_rich_para(doc, [
        ("Tipologi. ", True),
        ("(a) pidana oportunistik TBS; (b) potensi konflik tata kelola/relokasi; "
         "(c) non-Agrinas (PETI, tanah adat) — dipisahkan.", False),
    ])
    add_rich_para(doc, [
        ("Kualitas bukti. ", True),
        ("Kuat", True),
        (" untuk LP; hati-hati kerugian tipiring 2024 yang tidak masuk akal. ", False),
        ("Flag: ", True, True),
        ("time-series pencurian pra/pasca Agrinas; validasi klaster WJT (hijau vs tuntutan aktif).", False, True),
    ], space_after=8)

    # --- 4.6 Pelalawan ---
    add_heading_custom(doc, "4.6 Polres Pelalawan", level=2)
    add_rich_para(doc, [
        ("Ringkasan. ", True),
        ("9 kebun KSO seluruhnya hijau; konflik KSO resume 0. Dimensi terpisah: ", False),
        ("TNTN", True),
        (" (~81.793 Ha verifikasi; 7 desa; 5.733 KK) dengan intensitas aksi tinggi.", False),
    ])
    add_rich_para(doc, [
        ("Inventaris KSO. ", True),
        ("Mitra Unggul Pusaka, Sari Lembah Subur, Mekar Sari, Serikat Putra, Gandaerah (bagian), "
         "Eka Sari Lorena, Guna Dodos, Agrita Sari Prima, Viktorindo — mitra koperasi/poktan/Agrinas. "
         "Agregat lokal: disita badan usaha 5.642,65 Ha; KSO 5.392,83 Ha (10 KSO); "
         "angka “180.797.180 ha” ditolak.", False),
    ])
    add_rich_para(doc, [
        ("TNTN (23.02.2026). ", True),
        ("8 unjuk rasa; 3 penolakan portal; 5 pengusiran Satgas; "
         "penggerak AMMP/FKPM/Forum Tata Kelola/LMND/KOMARI/ACER; "
         "penegakan: 6 tersangka perusakan, 3 tersangka KSDAE; situasi relatif kondusif pasca penangkapan.", False),
    ])
    add_rich_para(doc, [
        ("Tipologi. ", True),
        ("KSO: rendah. TNTN: konflik kawasan hutan/negara–masyarakat (bukan tipikal penerima KSO).", False),
    ])
    add_rich_para(doc, [
        ("Kualitas bukti. ", True),
        ("KSO tipis; TNTN kuat. ", False),
        ("Flag: ", True, True),
        ("analisis terpisah TNTN vs KSO; minta daftar 10 KSO lengkap + LP bernomor.", False, True),
    ], space_after=8)

    # --- 4.7 Bengkalis ---
    add_heading_custom(doc, "4.7 Polres Bengkalis", level=2)
    add_rich_para(doc, [
        ("Ringkasan. ", True),
        ("8 kebun; resume 1 merah / 0 kuning / 7 hijau — ", False),
        ("tidak selaras", True),
        (" dengan kartu/Mei 2026 yang menampilkan beberapa kuning dan konflik berulang.", False),
    ])
    add_para(doc, "Hotspot:", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
    add_bullet(doc, "bentrok Des 2025 – Mei 2026; multi-LP; kerusakan kendaraan; "
               "panen liar ~150 Ha; pembakaran pos (Jan 2026).",
               bold_prefix="SIS – PAB (Mandau): ")
    add_bullet(doc, "penolakan/unjuk rasa KSO.",
               bold_prefix="CV Hendrik Padang & CV Sepakat Bersama Ali: ")
    add_bullet(doc, "mediasi gagal 28 Jul 2026, bentrok 29 Jul.",
               bold_prefix="Mutiara Naga – PKU/Agrinas: ")
    add_rich_para(doc, [
        ("Tipologi. ", True),
        ("Perebutan kelola sitaan antar eks-karyawan/KSO; penolakan penerima KSO; "
         "eskalasi berulang di lokasi sama.", False),
    ])
    add_rich_para(doc, [
        ("Kualitas bukti. ", True),
        ("Kuat. ", True),
        ("Flag: ", True, True),
        ("koreksi klaster Bengkalis; pantau Mutiara Naga sebagai kandidat naik klaster.", False, True),
    ], space_after=8)

    # --- 4.8 Kampar ---
    add_heading_custom(doc, "4.8 Polres Kampar", level=2)
    add_rich_para(doc, [
        ("Ringkasan. ", True),
        ("Intelkam hanya 4 kebun (3 kuning / 1 hijau), tetapi data lokal mencatat "
         "Agrinas 29.174,71 Ha (21 estate) dan LP/pengaduan agraria 2024–2026 paling terstruktur.", False),
    ])
    add_rich_para(doc, [
        ("Inventaris. ", True),
        ("Johan Sentosa (dikelola sendiri, hotspot Jun 2026); Torus Ganda Tambusan Timur (luas besar); "
         "PSPI, Sarindo/Kepau Jaya, CV Makmur Jaya Sentosa (kuning); Ciliandra (hijau); "
         "subset KSO ~9.571 Ha.", False),
    ])
    add_rich_para(doc, [
        ("Kasus Agrinas Jun 2026 (Pasir Sialang). ", True),
        ("LP/B/165, 168, 169, 190, 197/VI/2026 — pencurian/pengeroyokan terkait klaim ganda lahan. "
         "Agregat konflik agraria satker: 22 (2026), 27 (2025), 20 (2024) — mayoritas sengketa masyarakat "
         "yang tidak selalu terkait Agrinas.", False),
    ])
    add_rich_para(doc, [
        ("Tipologi. ", True),
        ("Klaim ganda di kebun Agrinas; penolakan Gapoktan terhadap KSO; "
         "PAM non-BUJP (CV Makmur — Flores).", False),
    ])
    add_rich_para(doc, [
        ("Kualitas bukti. ", True),
        ("Kuat", True),
        (" untuk LP; diskrepansi jumlah kebun Intelkam vs Disbun lokal perlu rekonsiliasi. ", False),
        ("Flag: ", True, True),
        ("pisahkan LP Agrinas vs LP agraria umum; rekonsiliasi 4 vs 21 estate.", False, True),
    ], space_after=8)

    # --- 4.9 Siak ---
    add_heading_custom(doc, "4.9 Polres Siak", level=2)
    add_rich_para(doc, [
        ("Ringkasan. ", True),
        ("6 kebun (semua hijau). Sinyal konflik lebih banyak dari rekap Setda 2024–2025 "
         "(batas kampung, HTI Arara, SSL, PTPN V) ketimbang KSO Agrinas sawit.", False),
    ])
    add_rich_para(doc, [
        ("Inventaris. ", True),
        ("Ivo Mas–Tiga Bintang Sinergi; BMI–Agung Anugerah Sawit; "
         "entitas hutan/HTI (Arara, RAPP, Tahura) dengan angka luas bermasalah.", False),
    ])
    add_rich_para(doc, [
        ("Tipologi. ", True),
        ("Administratif/izin/batas; satu kasus anarkis (bakar kantor PT SSL).", False),
    ])
    add_rich_para(doc, [
        ("Kualitas bukti. ", True),
        ("Tipis", True),
        (" untuk Agrinas–KSO; tabel luas tidak dapat dijumlah. ", False),
        ("Flag: ", True, True),
        ("filter hanya kebun sawit PKH–Agrinas; buang/verifikasi baris HTI.", False, True),
    ], space_after=8)

    # --- 4.10 Dumai ---
    add_heading_custom(doc, "4.10 Polres Dumai", level=2)
    add_rich_para(doc, [
        ("Ringkasan. ", True),
        ("3 kebun — semua kuning.", False),
    ])
    make_table(doc,
        ["Estate", "KSO/Pihak", "Isu"],
        [
            ["PT Duta Mas Makmur Perkasa", "PT Riden Jaya Konstruksi",
             "Bentrok tumpang tindih / take-over Feb 2026"],
            ["PT Sinar Riau Palm Oil", "self/KSO", "Perebutan pekerjaan"],
            ["PT Pelintung Jaya Bersama", "self", "Perebutan pekerjaan"],
        ],
        col_widths=[5.5, 5, 6])
    add_rich_para(doc, [
        ("Catatan. ", True),
        ("DMMP juga muncul di lembar Bengkalis (parcel lintas wilayah) — risiko double-count. ", False),
        ("Kualitas bukti: sedang. ", False),
        ("Flag: ", True, True),
        ("satu ID kebun multi-Polres; tarik LP Dumai bernomor.", False, True),
    ], space_after=8)

    # --- 4.11 Meranti ---
    add_heading_custom(doc, "4.11 Polres Kepulauan Meranti", level=2)
    add_rich_para(doc, [
        ("Ringkasan. ", True),
        ("Intelkam: 1 perusahaan (Sumatra Riang Lestari, hijau). Laporan lokal: NIHIL seluruh klaster "
         "Agrinas/PKH/poktan/masyarakat. File Polsek Tebing Tinggi corrupt.", False),
    ])
    add_rich_para(doc, [
        ("Kualitas bukti. ", True),
        ("Tipis / NIHIL. ", True),
        ("Flag: ", True, True),
        ("protocol gap-fill; kirim ulang template satker; jangan infer konflik dari ketiadaan data.", False, True),
    ], space_after=8)

    # --- 4.12 Pekanbaru ---
    add_heading_custom(doc, "4.12 Polresta Pekanbaru", level=2)
    add_rich_para(doc, [
        ("Ringkasan. ", True),
        ("Agrinas / sitaan PKH / KSO / potensi konflik: NIHIL. Terdapat 5 perusahaan sawit non-Agrinas "
         "dan 7 LP pencurian di Polsek Rumbai/Rumbai Barat 2025–2026 — pidana oportunistik masyarakat, "
         "bukan konflik tata kelola Agrinas–KSO.", False),
    ])
    add_rich_para(doc, [
        ("Kualitas bukti. ", True),
        ("Sedang", True),
        (" (jelas NIHIL untuk objek kajian). ", False),
        ("Flag: ", True, True),
        ("jaga pemisahan tipologi dalam dashboard agar tidak menggelembungkan “konflik Agrinas”.", False, True),
    ], space_after=8)

    doc.add_page_break()

    # ═══ 5. SINTESIS ═══
    add_heading_custom(doc, "5. SINTESIS LINTAS 12 POLRES", level=1)

    add_heading_custom(doc, "5.1 Peta Panas Komparatif", level=2)
    make_table(doc,
        ["Koridor", "Satker", "Karakter"],
        [
            ["Utara (eskalasi fisik)", "Rohul, Rohil, Bengkalis, Dumai",
             "Merah/kuning; PAM; bentrok; korban"],
            ["Selatan (volume & tuntutan kelola)", "Inhu, Kuansing, Inhil",
             "Banyak kebun; plasma/MHA/relokasi; pidana TBS"],
            ["Tengah (campuran)", "Kampar, Siak, Pelalawan",
             "Kampar LP bersih; Siak admin/HTI; Pelalawan TNTN"],
            ["Nihil Agrinas", "Pekanbaru, Kep. Meranti",
             "Monitoring ringan / gap-fill"],
        ],
        col_widths=[5, 5.5, 6])

    add_heading_custom(doc, "5.2 Rantai Tata Kelola Berisiko", level=2)
    add_code_block(doc,
        "Sita Satgas PKH → titip kelola Agrinas → penunjukan KSO\n"
        "    → (sering) penerima non-tempatan + PAM non-BUJP\n"
        "    → batas/plasma/MHA belum selesai\n"
        "    → gesekan panen → LP / bentrok / (Rohul) MD")
    add_para(doc,
        "Titik gagal paling berulang: penunjukan KSO tanpa legitimasi sosial lokal "
        "dan pengamanan non-formal.",
        first_indent=1.0)

    add_heading_custom(doc, "5.3 Struktural vs Oportunistik", level=2)
    make_table(doc,
        ["Jenis", "Ciri", "Contoh Satker"],
        [
            ["Struktural tata kelola",
             "Penolakan KSO, PAM vs PAM, klaim plasma/MHA, TNTN",
             "Rohul, Rohil, Bengkalis, Inhu, Pelalawan (TNTN), Kampar (Johan Sentosa)"],
            ["Pidana oportunistik di areal kebun",
             "Pencurian TBS berulang, tipiring, RJ",
             "Kuansing, Pekanbaru (non-Agrinas), sebagian Inhu/SWP"],
        ],
        col_widths=[4, 6, 6.5])
    add_para(doc,
        "Implikasi Harda: prioritas pemetaan konflik agraria pada rantai penunjukan & pengamanan KSO, "
        "bukan semata statistik pencurian TBS.",
        first_indent=1.0)

    add_heading_custom(doc, "5.4 Diskoneksi Data Harda vs Intelkam vs Reskrim", level=2)
    make_table(doc,
        ["Lapisan", "Kekuatan", "Kelemahan"],
        [
            ["Intelkam", "Cakupan 130 + klaster + bentrok",
             "Underestimate lokal (hijau yang bergesekan); inkonsistensi kartu vs resume"],
            ["Reskrim/LP satker", "Nomor LP, pasal, proses",
             "Tidak seragam; Rohil/Pelalawan tipis; Kuansing campur non-agraria"],
            ["Harda framing", "Tipologi yuridis–sosial–keamanan",
             "Belum punya master LP–kebun terpadu"],
            ["Disbun/Pemda (Siak/Pelalawan)", "Konteks izin/batas",
             "Angka luas sering tidak sinkron dengan Polri"],
        ],
        col_widths=[4.5, 5.5, 6.5])

    add_heading_custom(doc, "5.5 Temuan Komparatif Kunci", level=2)
    temuan = [
        "Korban jiwa/luka hampir seluruhnya di koridor utara terkait PAM/KSO.",
        "Jumlah kebun ≠ intensitas kekerasan (Inhu terbanyak kebun, Rohul paling mematikan).",
        "Klaster hijau tidak menjamin aman bila sumber satker mencatat tuntutan aktif "
        "(WJT, EX Palm, Mutiara Naga).",
        "TNTN Pelalawan adalah portofolio risiko terpisah dari KSO Agrinas.",
        "Double-count geografis (DMMP Bengkalis/Dumai; Gandaerah Pelalawan/Inhu; "
        "Torus naming Kampar/Rohul) harus dibersihkan sebelum skor risiko.",
    ]
    for i, t in enumerate(temuan, 1):
        add_para(doc, f"{i}. {t}", size=11, space_after=3, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_page_break()

    # ═══ 6. REKOMENDASI ═══
    add_heading_custom(doc, "6. REKOMENDASI ANALISIS YANG MUNGKIN DILAKUKAN", level=1)
    add_para(doc, "(Agenda riset/analitik — bukan perintah operasi)",
             italic=True, size=10, color=GRAY, align=WD_ALIGN_PARAGRAPH.LEFT)

    rekomendasi = [
        ("R1. Matriks Risiko Satker", "Tinggi", [
            "Tujuan: skor komposit per Polres/kebun.",
            "Input: jml kebun, % belum KSO, klaster, frekuensi LP, korban, PAM non-BUJP.",
            "Output: ranking bulanan Unit II Harda.",
            "Sumber siap: sheet ringkasan_polres + kebun.",
        ]),
        ("R2. Analisis Jaringan Aktor KSO", "Tinggi", [
            "Tujuan: deteksi penerima/PAM yang muncul multi-lokasi.",
            "Contoh seed: Bernas Mulya Mandiri, PT Runggu, PAB, Majuma, Digjaya/UTS, Riden Jaya.",
            "Output: graf aktor–kebun–Polres; daftar “hub risiko”.",
        ]),
        ("R3. Validasi Klaster Intelkam vs Kejadian Aktual", "Tinggi", [
            "Tujuan: uji apakah merah/kuning/hijau masih valid per Jul 2026.",
            "Metode: cocokkan kartu klaster dengan LP/bentrok 2025–2026.",
            "Kandidat naik klaster: Mutiara Naga (Bengkalis), WJT/Pesikaian (Kuansing), "
            "EX Palm (Inhu), Johan Sentosa (Kampar).",
        ]),
        ("R4. Timeline Eskalasi PKH → KSO → LP/Bentrok", "Tinggi", [
            "Tujuan: pola lead-time dari penunjukan KSO hingga gesekan.",
            "Pilot: Rohul (Majuma), Bengkalis (PAB–SIS), Rohil (UTS), Dumai (DMMP).",
            "Output: diagram waktu + window kritis pengawasan.",
        ]),
        ("R5. Tipologi Hukum Harda", "Sedang", [
            "Tujuan: klasifikasi tiap kasus: pidana murni / sengketa perdata / "
            "administratif (HGU–SIUP–TNTN) / campuran.",
            "Manfaat: alokasi ke Reskrim vs binmas/mediasi vs lintas sektoral.",
        ]),
        ("R6. Spasial Kasar Kecamatan/Desa", "Sedang", [
            "Tujuan: hotspot TKP dari lokasi yang sudah ada di matriks.",
            "Syarat: geocoding manual desa (Bonai, Mandau, Pasir Sialang, Pesikaian, Pelintung, dll.).",
            "Output: peta panas sederhana untuk paparan.",
        ]),
        ("R7. Gap-Filling Protocol Satker Tipis", "Sedang", [
            "Target: Meranti, Pekanbaru (konfirmasi berkala), Rohil (itemisasi LP), "
            "Pelalawan (10 KSO + LP), Siak (filter sawit PKH).",
            "Output: template isian seragam = kolom sheet kebun + kasus_lp.",
        ]),
        ("R8. Dashboard Monitoring", "Rendah (setelah R1–R3)", [
            "Tujuan: pantauan bulanan klaster, LP baru, status KSO.",
            "Syarat: matriks Excel distabilkan dan diisi ulang oleh satker.",
        ]),
    ]

    for judul, prioritas, points in rekomendasi:
        add_heading_custom(doc, f"{judul}  —  Prioritas: {prioritas}", level=2)
        for pt in points:
            add_bullet(doc, pt)

    add_heading_custom(doc, "Lima Prioritas Segera untuk Pimpinan Unit II", level=2)
    prioritas_segera = [
        "R1 skor risiko satker",
        "R3 validasi klaster (utamanya Bengkalis & kebun hijau bergesekan)",
        "R2 jaringan KSO/PAM",
        "R4 timeline eskalasi 4 kasus merah/kuning utama",
        "R7 gap-fill Rohil LP + Pelalawan daftar KSO + Meranti template",
    ]
    for i, p in enumerate(prioritas_segera, 1):
        add_para(doc, f"{i}. {p}", size=11, space_after=3, align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_page_break()

    # ═══ 7. LAMPIRAN ═══
    add_heading_custom(doc, "7. LAMPIRAN SUMBER", level=1)
    sumber = [
        "12.02.2026 - PERMASALAHAN KSO PT. AGRINAS DI WILKUM POLDA RIAU.pdf",
        "23.02.2026 - PERKEMBANGAN SITUASI TNTN DAN KSO PT. AGRINAS.pdf",
        "PERMASALAHAN TANAH HARDA RIAU 2026 Sudah di Ganti.pdf",
        "LUASAN PERKEBUNAN SAWIT KESELURUHAN DI WILAYAH KAB. INHU (3).docx",
        "POTENSI KONFLIK AGRARIA 2024 - 2026 POLDA RIAU (1).docx (isi: Kampar)",
        "POTENSI KONFLIK LAHAN PT AGRINAS TAHUN 2026.docx",
        "DATA LUAS PERKEBUNAN KAB. KAMPAR.docx",
        "POLRES ROKAN HULU.docx",
        "DATA KONFLIK PERKEBUNAN DI WILAYAH HUKUM KAB. HILIR.pdf",
        "DATA LAHAN KEBUN SAWIT SITA PKH BULAN MEI 2026.pdf",
        "DATA PERKEBUNAN BERDASARKAN STATUS KSO.pdf",
        "POTENSI KONFLIK 2026.pdf",
        "POLRES KUANSING 2020-2026.pdf",
        "LUASAN KEBUN SAWIT DI WILKUM KUANSING.pdf",
        "DATA PERKEBUNAN KELAPA SAWIT WILKUM POLRES PELALAWAN.pdf",
        "DATA KSO DAN LUAS LAHAN WILKUM SIAK.pdf",
        "Bismillah.docx",
        "DATA TINDAK PIDANA PENCURIAN SAWIT TAHUN 2025 DAN 2026.docx",
        "LUASAN PERKEBUNAN SAWIT KESELURUHAN DI WILAYAH KAB. KEP. MERANTI.docx",
        "LUASAN ... POLSEK TEBING TINGGI.docx (corrupt)",
        "DATA LAHAN ... MEI 2026 (1).pdf (duplikat)",
    ]
    make_table(doc,
        ["No", "File Sumber"],
        [[str(i), s] for i, s in enumerate(sumber, 1)],
        col_widths=[1.5, 15])

    doc.add_paragraph()
    add_para(doc,
        "Dokumen ini disusun untuk kepentingan analisis Unit II Harda Ditreskrimum Polda Riau "
        "berdasarkan berkas yang tersedia di workspace pada tanggal penyusunan (3 Agustus 2026). "
        "Sifat dokumen: deskripsi–sintesis; bukan verifikasi lapangan dan bukan perintah operasi.",
        size=10, italic=True, color=GRAY, space_before=12)

    # tanda tangan
    doc.add_paragraph()
    add_para(doc, "Mengetahui / Menyusun,", size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_para(doc, "Unit II Harda", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_para(doc, "Ditreskrimum Polda Riau", size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=36)
    add_para(doc, "________________________", size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_para(doc, "Analis Unit II Harda", size=10, italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT, color=GRAY)

    doc.save(OUT)
    print(f"OK: {OUT}")


if __name__ == "__main__":
    build()
