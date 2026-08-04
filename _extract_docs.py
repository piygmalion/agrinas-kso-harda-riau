# -*- coding: utf-8 -*-
import os
from docx import Document

base = r"C:\Users\Patron\Downloads\sawit lagi"
out = os.path.join(base, "_extract")
os.makedirs(out, exist_ok=True)

docx_files = [f for f in os.listdir(base) if f.lower().endswith(".docx")]
for f in docx_files:
    path = os.path.join(base, f)
    try:
        doc = Document(path)
        lines = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                lines.append(t)
        for ti, table in enumerate(doc.tables):
            lines.append(f"\n=== TABLE {ti+1} ===")
            for row in table.rows:
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                lines.append(" | ".join(cells))
        text = "\n".join(lines)
        with open(os.path.join(out, f + ".txt"), "w", encoding="utf-8") as fh:
            fh.write(text)
        print(f"OK {f}: {len(text)} chars, {len(doc.tables)} tables")
    except Exception as e:
        print(f"FAIL {f}: {e}")
